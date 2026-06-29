#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Module pour convertir des fichiers DOCX en HTML
Utilisé par le générateur de PDF
"""

import os
import tempfile
import shutil
import subprocess
from pathlib import Path
import traceback

def docx_to_html(docx_path, output_dir=None):
    """
    Convertit un fichier DOCX en HTML
    
    Args:
        docx_path (str): Chemin vers le fichier DOCX
        output_dir (str): Répertoire de sortie pour le HTML
        
    Returns:
        str: Chemin vers le fichier HTML généré
    """
    try:
        # Vérifier si le fichier existe
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Le fichier {docx_path} n'existe pas")
        
        # Créer un répertoire temporaire si aucun n'est spécifié
        if output_dir is None:
            output_dir = tempfile.mkdtemp()
        else:
            os.makedirs(output_dir, exist_ok=True)
        
        # Nom de fichier de sortie
        base_name = os.path.basename(docx_path)
        name_without_ext = os.path.splitext(base_name)[0]
        html_path = os.path.join(output_dir, f"{name_without_ext}.html")
        
        print(f"Conversion du fichier DOCX {docx_path} en HTML...")
        
        # Tenter de convertir avec pandoc (s'il est installé)
        try:
            cmd = ["pandoc", docx_path, "-o", html_path]
            result = subprocess.run(cmd, capture_output=True, text=True, check=True)
            print("✅ Conversion réussie avec pandoc")
            return html_path
        except (subprocess.CalledProcessError, FileNotFoundError) as e:
            print(f"Pandoc n'est pas disponible ou a échoué: {e}")
            print("Essai de méthodes alternatives...")
        
        # Méthode alternative: utiliser python-docx et convertir manuellement
        try:
            from docx import Document
            
            # Charger le document
            doc = Document(docx_path)
            
            # Convertir en HTML basique
            html_content = ["<!DOCTYPE html>", "<html>", "<head>", 
                          "<meta charset='utf-8'>", 
                          "<title>Convocation</title>",
                          "<style>",
                          "body { font-family: Arial, sans-serif; }",
                          "table { border-collapse: collapse; width: 100%; }",
                          "table, th, td { border: 1px solid #ddd; }",
                          "th, td { padding: 8px; text-align: left; }",
                          "</style>",
                          "</head>", 
                          "<body>"]
            
            # Extraire le texte de chaque paragraphe
            for para in doc.paragraphs:
                # Ignorer les paragraphes vides
                if para.text.strip():
                    style = ""
                    if para.style.name.startswith('Heading'):
                        level = para.style.name[-1]  # Obtenir le niveau du titre (1, 2, etc.)
                        html_content.append(f"<h{level}>{para.text}</h{level}>")
                    else:
                        html_content.append(f"<p>{para.text}</p>")
            
            # Extraire les tableaux
            for table in doc.tables:
                html_content.append("<table>")
                for row in table.rows:
                    html_content.append("<tr>")
                    for cell in row.cells:
                        # Obtenir tout le texte de la cellule
                        cell_text = ""
                        for paragraph in cell.paragraphs:
                            cell_text += paragraph.text + " "
                        html_content.append(f"<td>{cell_text.strip()}</td>")
                    html_content.append("</tr>")
                html_content.append("</table>")
            
            html_content.append("</body></html>")
            
            # Écrire le contenu HTML dans un fichier
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write("\n".join(html_content))
            
            print("✅ Conversion réussie avec python-docx")
            return html_path
            
        except Exception as e:
            print(f"La conversion avec python-docx a échoué: {e}")
            print(f"Détails: {traceback.format_exc()}")
            
            # Dernière solution: créer un HTML de base avec un message d'erreur
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write("""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Erreur de Conversion</title>
</head>
<body>
    <h1>Erreur de Conversion du Template</h1>
    <p>Le fichier DOCX n'a pas pu être converti en HTML.</p>
    <p>Veuillez utiliser un template au format HTML ou installer Pandoc pour la conversion automatique.</p>
</body>
</html>""")
            
            print("⚠️ Création d'un HTML de base avec message d'erreur")
            return html_path
    
    except Exception as e:
        print(f"Erreur lors de la conversion DOCX -> HTML: {e}")
        print(f"Détails: {traceback.format_exc()}")
        raise

if __name__ == "__main__":
    # Test de la fonction
    import sys
    
    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
        output_dir = "templates" if len(sys.argv) <= 2 else sys.argv[2]
        
        try:
            html_path = docx_to_html(docx_path, output_dir)
            print(f"Fichier HTML généré: {html_path}")
        except Exception as e:
            print(f"Erreur: {e}")
    else:
        print("Usage: python docx_to_html.py chemin_fichier_docx [repertoire_sortie]")